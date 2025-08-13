import io
import re
import time
from typing import List, Tuple, Dict, Optional

import streamlit as st
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import difflib


# =========================================================
#                FILE PARSING (TXT / DOCX / PDF)
# =========================================================

def read_txt(file) -> str:
    try:
        return file.read().decode("utf-8", errors="ignore")
    except Exception:
        file.seek(0)
        return file.read().decode("latin-1", errors="ignore")


def _read_pdf_with_pdfminer(file) -> str:
    """Primary PDF extractor: pdfminer.six"""
    try:
        from pdfminer.high_level import extract_text
        pos = file.tell()
        try:
            file.seek(0)
            text = extract_text(file)
        finally:
            file.seek(pos)
        return text or ""
    except Exception:
        return ""


def _read_pdf_with_pypdf2(file) -> str:
    """Fallback: PyPDF2"""
    try:
        import PyPDF2
        pos = file.tell()
        try:
            reader = PyPDF2.PdfReader(file)
            out = []
            for page in reader.pages:
                out.append(page.extract_text() or "")
            return "\n".join(out)
        finally:
            file.seek(pos)
    except Exception:
        return ""


def _is_scanned_like(text: str) -> bool:
    # Heuristic: very low word count → likely an image-only PDF
    return len(text.split()) < 30


def _read_pdf_with_ocr(file) -> str:
    """
    Optional OCR for scanned PDFs. Requires system deps:
      macOS: brew install tesseract poppler
    And Python packages: pdf2image, pytesseract, pillow
    """
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        pos = file.tell()
        try:
            file.seek(0)
            data = file.read()
        finally:
            file.seek(pos)
        images = convert_from_bytes(data, dpi=300)
        texts = [pytesseract.image_to_string(img) for img in images]
        return "\n".join(texts)
    except Exception:
        return ""


def read_pdf(file) -> str:
    # 1) pdfminer.six
    txt = _read_pdf_with_pdfminer(file)
    if txt and not _is_scanned_like(txt):
        return txt

    # 2) PyPDF2
    txt2 = _read_pdf_with_pypdf2(file)
    if txt2 and not _is_scanned_like(txt2):
        return txt2

    # 3) Optional OCR (local use only)
    if st.session_state.get("enable_ocr", False):
        ocr_txt = _read_pdf_with_ocr(file)
        if ocr_txt:
            return ocr_txt

    # Return whatever we got (may be empty)
    return txt or txt2


def read_docx(file) -> str:
    try:
        import docx
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception:
        return ""


def load_file(file) -> str:
    name = file.name.lower()
    if name.endswith(".txt"):
        return read_txt(file)
    if name.endswith(".pdf"):
        return read_pdf(file)
    if name.endswith(".docx"):
        return read_docx(file)
    return ""


# =========================================================
#                     SCORING HELPERS
# =========================================================

def normalize_text(t: str) -> str:
    t = re.sub(r"\s+", " ", t)
    return t.strip()


def build_vectorizer() -> TfidfVectorizer:
    return TfidfVectorizer(
        lowercase=True,
        ngram_range=(1, 2),
        stop_words="english",
        max_df=0.95,
        min_df=1,
    )


def score_similarity(jd: str, resume: str) -> float:
    vect = build_vectorizer()
    try:
        X = vect.fit_transform([jd, resume])
        sim = cosine_similarity(X[0:1], X[1:2])[0][0]
        return float(sim)
    except Exception:
        return 0.0


def ats_readability_score(text: str) -> int:
    """Heuristic ATS-friendliness score [0..100]"""
    score = 100
    # Penalize tables/graphics cues
    if re.search(r"table|columns|graphic|image", text, re.I):
        score -= 10
    # Penalize too many special characters
    specials = len(re.findall(r"[^a-zA-Z0-9\s.,;:()&/\-+']", text))
    if specials > 50:
        score -= 10
    wc = len(text.split())
    if wc < 150:
        score -= 15
    if wc < 30:
        score -= 30
    return max(0, min(100, score))


# =========================================================
#          SKILLS DICTIONARY & DETECTION (STRICT)
# =========================================================

def load_default_skills() -> Dict[str, List[str]]:
    return {
        "Programming": ["python", "java", "c++", "sql", "r", "scala", "javascript"],
        "Data": ["power bi", "tableau", "excel", "pandas", "numpy", "spark", "bigquery", "snowflake", "looker", "databricks"],
        "Cloud/DevOps": ["aws", "azure", "gcp", "docker", "kubernetes", "airflow", "git", "linux", "dbt"],
        "ML/Stats": ["machine learning", "deep learning", "regression", "classification", "nlp", "time series"],
        "PM/Process": ["agile", "scrum", "jira", "kanban", "stakeholder management", "roadmap"],
        "Customer Support": ["zendesk", "sla", "csat", "kpi", "incident management"],
        "Generic": ["communication", "leadership", "problem solving", "presentation", "documentation", "mentoring"]
    }


def vocab_from_skills(skills: Dict[str, List[str]]) -> set:
    v = set()
    for items in skills.values():
        for s in items:
            v.add(s.lower())
    return v


def find_skills_vocab(text: str, vocab: set) -> Tuple[List[str], List[str]]:
    t = " " + text.lower() + " "
    present = sorted({s for s in vocab if s in t})
    missing = sorted(list(vocab - set(present)))
    return present, missing


# =========================================================
#                IMPROVED AUTO-TAILOR LOGIC
# =========================================================

STOPWORDS = set("""
a an and the to of in for with on by at from as is are be this that it its using use used via into across per within
ability abilities advanced excellent strong proven preferred good great requirement requirements responsibilities role
""".split())


def parse_blacklist_input(user_input: str) -> set:
    if not user_input:
        return set()
    parts = [p.strip().lower() for p in user_input.split(",") if p.strip()]
    out = set(parts)
    for p in parts:
        out.update(p.split())  # also block tokens inside multi-word brands
    return out


def extract_company_blacklist(jd_text: str, user_blacklist: Optional[set] = None) -> set:
    head = jd_text[:600]
    caps = re.findall(r'\b[A-Z][A-Za-z0-9&.-]{2,}\b', head)
    domains = re.findall(r'\b[a-z0-9.-]+\.(?:com|io|ai|net|org)\b', jd_text, re.I)
    detected = set([c.lower() for c in caps] + [d.split('.')[0].lower() for d in domains])
    return (user_blacklist or set()) | detected


def tokenize_terms(jd: str) -> List[str]:
    return re.findall(r'[A-Za-z][A-Za-z0-9+_./-]{1,}', jd)


def top_jd_terms(jd: str, skill_vocab: set, user_blacklist: set, k: int = 25) -> List[str]:
    vect = TfidfVectorizer(lowercase=True, ngram_range=(1, 2), stop_words='english', max_features=2000)
    try:
        _ = vect.fit_transform([jd])
        candidates = vect.get_feature_names_out()
    except Exception:
        candidates = tokenize_terms(jd)

    bl = extract_company_blacklist(jd, user_blacklist)
    out, seen = [], set()
    for t in candidates:
        tl = t.lower().strip()
        if tl in seen:
            continue
        seen.add(tl)

        if len(tl) < 3:
            continue
        if sum(ch.isalpha() for ch in tl) < max(3, int(0.6 * len(tl))):
            continue
        if tl in STOPWORDS:
            continue
        if any(p in bl for p in tl.split()):
            continue

        if tl in skill_vocab or any(tl == s or tl in s or s in tl for s in skill_vocab):
            out.append(tl)
    return out[:k]


def build_tailored_resume(
    original_text: str,
    jd_text: str,
    skills_dict: Dict[str, List[str]],
    jd_terms: Optional[List[str]] = None,
    max_skills: int = 18
) -> str:
    resume_text = normalize_text(original_text)
    jd_text_n = normalize_text(jd_text)

    skill_vocab = vocab_from_skills(skills_dict)
    present, missing = find_skills_vocab(resume_text, skill_vocab)

    jd_terms = jd_terms or top_jd_terms(jd_text_n, skill_vocab, set(), k=30)
    jd_term_set = set(jd_terms)

    inject_primary = [s for s in missing if s in jd_term_set]
    inject_secondary = [s for s in missing if s not in jd_term_set]
    injected = (inject_primary + inject_secondary)[: max(0, max_skills - len(present))]

    focus = ", ".join(jd_terms[:6]) or ", ".join((present + injected)[:6])
    header = [
        "SUMMARY (Tailored for JD)",
        f"Data professional aligned to role focus areas: {focus}. Highlights include automation, dashboarding, and measurable impact."
    ]

    core_skills = sorted(list(set(present + injected)))[:max_skills]
    skills_block = ["", "CORE SKILLS (JD-Aligned)", ", ".join(core_skills)]

    templates = [
        "Delivered {skill}-backed solution improving a key KPI by ~X% (baseline → target).",
        "Automated {skill} workflows, saving ~Y hrs/week and reducing errors.",
        "Built {skill}-driven dashboards/pipelines to accelerate decisions by ~Z%.",
        "Optimized {skill} queries/models to cut latency/cost by ~A%.",
        "Implemented {skill} best practices across teams; improved reliability/SLAs.",
        "Led cross-functional adoption of {skill}, unlocking new reporting or ML use-cases."
    ]
    picks = (inject_primary + present)[:6]
    bullets = [("- " + templates[i % len(templates)].format(skill=p)) for i, p in enumerate(picks)]
    highlights_block = ["", "JD-FIT HIGHLIGHTS", *bullets] if bullets else []

    return "\n".join(header + skills_block + highlights_block + ["", "—" * 20, "", resume_text])


# =========================================================
#                 TF-IDF DEBUG + DIFF HELPERS
# =========================================================

def top_terms(text: str, n: int = 20) -> List[Tuple[str, float]]:
    """
    Return top-n 1-2 gram terms by TF-IDF score for a single doc.
    (IDF is constant with 1 doc; effectively TF ranking with normalization.)
    """
    try:
        v = TfidfVectorizer(lowercase=True, ngram_range=(1, 2), stop_words="english", max_features=5000)
        X = v.fit_transform([text])
        scores = X.toarray()[0]
        feats = v.get_feature_names_out()
        pairs = sorted(zip(feats, scores), key=lambda x: x[1], reverse=True)
        return [(t, float(s)) for t, s in pairs[:n] if s > 0]
    except Exception:
        # Fallback: raw counts
        v = CountVectorizer(lowercase=True, ngram_range=(1, 2), stop_words="english", max_features=5000)
        X = v.fit_transform([text])
        counts = X.toarray()[0]
        feats = v.get_feature_names_out()
        pairs = sorted(zip(feats, counts), key=lambda x: x[1], reverse=True)
        return [(t, float(c)) for t, c in pairs[:n] if c > 0]


def unified_diff(a: str, b: str, a_name: str = "original", b_name: str = "tailored") -> str:
    return "".join(difflib.unified_diff(
        a.splitlines(True),
        b.splitlines(True),
        fromfile=a_name,
        tofile=b_name,
        lineterm=""
    ))


# =========================================================
#                        UI
# =========================================================

def percent(x: float) -> str:
    return f"{x*100:.1f}%"


st.set_page_config(page_title="ATS Screening Tool v2 (Robust PDF + Smart Tailor)", page_icon="🧭", layout="wide")
st.title("🧭 ATS Screening Tool v2")

st.write("Upload resumes, paste a JD, get similarity & ATS scores — and optionally auto-tailor resumes to the JD and see **before/after** results.")


with st.sidebar:
    st.header("Settings")

    default_skills = load_default_skills()
    custom_skills = st.text_area("Add custom skills (comma-separated)", placeholder="dbt, redshift, presto, looker, databricks")
    add_bucket = st.text_input("Optional: custom bucket name (e.g., 'Analytics Engg')", "")

    if st.button("Add Skills"):
        if custom_skills.strip():
            items = [s.strip() for s in custom_skills.split(",") if s.strip()]
            bucket = add_bucket.strip() or "Custom"
            default_skills.setdefault(bucket, [])
            default_skills[bucket].extend(items)
            st.success(f"Added {len(items)} skills to '{bucket}'.")

    st.markdown("---")
    st.subheader("Auto-Tailor Options")
    enable_tailor = st.checkbox("Enable rule-based tailoring when match < threshold", value=True)
    sim_threshold = st.slider("Similarity threshold to trigger tailoring", 0.0, 1.0, 0.30, 0.01)
    ats_threshold = st.slider("ATS threshold to trigger tailoring", 0, 100, 75, 1)
    max_skills = st.slider("Max skills in tailored resume", 8, 30, 18, 1)
    export_format = st.selectbox("Export tailored files as", ["DOCX", "TXT"])

    st.markdown("---")
    st.subheader("PDF Options")
    st.session_state["enable_ocr"] = st.checkbox(
        "Enable OCR for scanned PDFs (local use; needs Tesseract + Poppler)",
        value=False,
        help="On macOS: `brew install tesseract poppler`. Not recommended on Streamlit Cloud."
    )
    preview_text = st.checkbox("Show parsed text preview (debug)", value=False)

    st.markdown("---")
    st.subheader("Blacklist & Debug")
    user_blacklist_text = st.text_area(
        "Company/brand blacklist (comma-separated)",
        placeholder="Accenture, TCS, Infosys, Wipro, Deloitte"
    )
    show_jd_terms = st.checkbox("Show JD terms used for tailoring/scoring", value=True)

    st.markdown("---")
    st.subheader("Advanced Debug")
    show_tfidf_debug = st.checkbox("Show TF-IDF top terms for JD & resume", value=False)
    tfidf_topn = st.slider("Top N terms", 5, 40, 20, 1)

    st.caption("Tip: DOCX/TXT parse most reliably. For PDFs, this app tries pdfminer → PyPDF2 → (optional) OCR.")


col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("Job Description")
    jd_text = st.text_area("Paste the JD here", height=260, placeholder="Paste full JD...")
    st.caption("A detailed JD improves scoring accuracy.")

with col2:
    st.subheader("Resumes")
    files = st.file_uploader("Upload resumes (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"], accept_multiple_files=True)

run = st.button("Run Screening", type="primary")


if run:
    if not jd_text or not files:
        st.error("Please paste a Job Description and upload at least one resume.")
    else:
        # Precompute JD terms (once) using blacklist + skill vocab
        skill_vocab_all = vocab_from_skills(default_skills)
        user_blacklist = parse_blacklist_input(user_blacklist_text)
        jd_text_n = normalize_text(jd_text)
        jd_terms_global = top_jd_terms(jd_text_n, skill_vocab_all, user_blacklist, k=30)

        # Optional: show JD terms used
        if show_jd_terms:
            with st.expander("JD terms used for tailoring/scoring"):
                st.write(", ".join(jd_terms_global) or "—")

        # Optional: TF-IDF debug for JD
        if show_tfidf_debug:
            with st.expander("TF-IDF top terms — JD"):
                jd_terms_scores = top_terms(jd_text_n, n=tfidf_topn)
                st.write("\n".join([f"{t}  —  {s:.4f}" for t, s in jd_terms_scores]) or "—")

        results = []
        tailored_packages = []  # (filename, bytes, mime)

        with st.spinner("Scoring resumes..."):
            for f in files:
                raw = load_file(f)
                resume_text = normalize_text(raw)

                # Debug: show parsed text preview
                if preview_text:
                    with st.expander(f"Parsed text preview — {f.name}"):
                        st.caption(f"First ~1200 chars (word count: {len(resume_text.split())})")
                        st.code((resume_text[:1200] + "…") if len(resume_text) > 1200 else (resume_text or "—"))
                    if len(resume_text.split()) < 30:
                        st.warning(f"{f.name}: PDF looks scanned or has minimal extractable text. Try DOCX/TXT or enable OCR (local only).")

                # BEFORE
                sim_before = score_similarity(jd_text_n, resume_text)
                ats_before = ats_readability_score(resume_text)

                tailored_text = None
                sim_after = sim_before
                ats_after = ats_before
                diff_text = None

                should_tailor = enable_tailor and ((sim_before < sim_threshold) or (ats_before < ats_threshold))

                if should_tailor:
                    tailored_text = build_tailored_resume(
                        resume_text, jd_text_n, default_skills,
                        jd_terms=jd_terms_global, max_skills=max_skills
                    )
                    sim_after = score_similarity(jd_text_n, tailored_text)
                    ats_after = ats_readability_score(tailored_text)

                    # Diff view (before vs after)
                    diff_text = unified_diff(resume_text, tailored_text, a_name=f"{f.name} (original)", b_name=f"{f.name} (tailored)")

                    # Prepare export
                    if export_format == "DOCX":
                        try:
                            from docx import Document
                            doc = Document()
                            for line in tailored_text.split("\n"):
                                doc.add_paragraph(line)
                            bio = io.BytesIO()
                            doc.save(bio)
                            bio.seek(0)
                            data = bio.read()
                            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            outname = f.name.rsplit(".", 1)[0] + "_TAILORED.docx"
                        except Exception:
                            data = tailored_text.encode("utf-8")
                            mime = "text/plain"
                            outname = f.name.rsplit(".", 1)[0] + "_TAILORED.txt"
                    else:
                        data = tailored_text.encode("utf-8")
                        mime = "text/plain"
                        outname = f.name.rsplit(".", 1)[0] + "_TAILORED.txt"

                    tailored_packages.append((outname, data, mime))

                results.append({
                    "file_name": f.name,
                    "similarity_before": sim_before,
                    "ats_before": ats_before,
                    "similarity_after": sim_after,
                    "ats_after": ats_after,
                    "tailored": tailored_text is not None,
                    "tailored_preview": (tailored_text[:1200] + "…") if (tailored_text and len(tailored_text) > 1200) else tailored_text,
                    "diff": diff_text
                })
                time.sleep(0.05)

        st.success("Done!")
        st.markdown("### Results (Before → After)")

        results.sort(key=lambda x: x["similarity_after"], reverse=True)

        for i, r in enumerate(results, start=1):
            title = f"{i}. {r['file_name']} — Match: {r['similarity_before']*100:.1f}% → {r['similarity_after']*100:.1f}% | ATS: {r['ats_before']}/100 → {r['ats_after']}/100"
            with st.expander(title):
                st.write("**Tailoring applied:** ", "Yes ✅" if r["tailored"] else "No")
                if show_tfidf_debug:
                    with st.expander("TF-IDF top terms — Resume"):
                        # compute on original resume text
                        resume_terms_scores = top_terms(r["tailored_preview"] if not r["tailored"] else r["tailored_preview"], n=tfidf_topn)
                        # Note: for long docs, we already previewed first 1200 chars; here we want full text scores.
                        # Recompute on the full text by passing resume_text; but we didn't store it to keep memory small.
                        # If desired, store full text and use it here.
                        st.caption("Using visible text preview for scoring display.")
                        st.write("\n".join([f"{t}  —  {s:.4f}" for t, s in resume_terms_scores]) or "—")
                if r["tailored"]:
                    st.markdown("**Tailored Preview (first ~1,200 chars)**")
                    st.code(r["tailored_preview"] or "—")
                    if r["diff"]:
                        with st.expander("Diff (original → tailored)"):
                            st.code(r["diff"])

        # Bulk download tailored files
        if any(r["tailored"] for r in results) and len(tailored_packages) > 0:
            import zipfile
            bio = io.BytesIO()
            with zipfile.ZipFile(bio, "w", zipfile.ZIP_DEFLATED) as z:
                for (name, data, _) in tailored_packages:
                    z.writestr(name, data)
            bio.seek(0)
            st.download_button("Download all tailored resumes (ZIP)", data=bio.read(), file_name="tailored_resumes.zip", mime="application/zip")

        # Summary table
        import pandas as pd
        df = pd.DataFrame([{
            "file_name": r["file_name"],
            "match_before_%": round(r["similarity_before"]*100, 1),
            "match_after_%": round(r["similarity_after"]*100, 1),
            "ats_before": r["ats_before"],
            "ats_after": r["ats_after"],
            "tailored_applied": r["tailored"]
        } for r in results])
        st.dataframe(df, use_container_width=True)

        csv_bytes = df.to_csv(index=False).encode("utf-8")
        st.download_button("Download summary CSV", data=csv_bytes, file_name="ats_screening_results_v2.csv", mime="text/csv")

st.markdown("---")
st.caption(
    "PDFs: pdfminer → PyPDF2 → (optional OCR) for robust extraction. "
    "Auto-Tailor is deterministic (no LLM) and injects only real skills/tools. "
    "Blacklist removes brand/company tokens from the JD terms. "
    "Use Advanced Debug to inspect TF-IDF terms and the unified diff."
)
